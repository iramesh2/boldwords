import streamlit as st

# Must set page config first before any other Streamlit commands
st.set_page_config(page_title="Bold Text Extractor", layout="wide")

import os
import tempfile
import openai
from docx import Document
import re
from pathlib import Path
import json
import traceback

# Add debug mode
debug_mode = True

def debug_print(message):
    """Print debug messages if debug mode is enabled"""
    if debug_mode:
        st.write(f"🔍 DEBUG: {message}")

# Get the OpenAI API key from environment variable and set it
api_key = os.environ.get("API")
if api_key:
    debug_print(f"API key found: {api_key[:4]}...{api_key[-4:]}")
    openai.api_key = api_key
else:
    st.error("OpenAI API key not found. Please set the API environment variable.")
    debug_print("API key not found in environment variables")

def convert_docx_to_raw_text(docx_file):
    """Convert a docx file to a plain text for OpenAI processing"""
    debug_print(f"Converting document: {docx_file}")
    doc = Document(docx_file)
    text = []
    paragraph_count = 0
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
            paragraph_count += 1
    
    debug_print(f"Processed {paragraph_count} paragraphs")
    result = "\n".join(text)
    if debug_mode:
        st.expander("Document Text Preview").text(result[:500] + "..." if len(result) > 500 else result)
    return result

def identify_headers_with_openai(text):
    """Use OpenAI to identify potential section headers in the document"""
    try:
        debug_print("Starting header identification with OpenAI")
        if not openai.api_key:
            st.error("OpenAI API key is not set. Cannot identify headers.")
            return []
        
        debug_print(f"Using model: gpt-4o")
        debug_print(f"Document length: {len(text)} characters")
        
        # For debugging, show what model versions are available
        try:
            debug_print("OpenAI module version: " + openai.__version__)
        except:
            debug_print("Could not determine OpenAI module version")
            
        # Show what request we're making
        prompt = "Below is the text of a document. Please identify the main section headers that divide this document into logical parts. Return ONLY a JSON array of strings containing ONLY the header text."
        debug_print(f"Prompt: {prompt}")
        
        try:
            # First try the newer format
            debug_print("Attempting API call with newer ChatCompletion format")
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a document analysis assistant. Your task is to identify section headers in the document."},
                    {"role": "user", "content": f"{prompt} Example: [\"Introduction\", \"Methods\", \"Results\", \"Discussion\"]. Don't include any explanations, just the JSON array.\n\n{text}"}
                ],
                temperature=0.3,
            )
        except Exception as e1:
            debug_print(f"Error with ChatCompletion: {str(e1)}")
            try:
                # Try older format as fallback
                debug_print("Attempting API call with older completion format")
                response = openai.Completion.create(
                    engine="gpt-4o",
                    prompt=f"Identify the main section headers in this document. Return only a JSON array of headers:\n\n{text}",
                    max_tokens=1000,
                    temperature=0.3,
                )
                # Create compatible response format
                response = {"choices": [{"message": {"content": response.choices[0].text}}]}
                debug_print("Successfully used older format")
            except Exception as e2:
                debug_print(f"Error with Completion fallback: {str(e2)}")
                raise
        
        debug_print("OpenAI request successful")
        
        # Handle different response formats
        try:
            # Try newer format first
            headers_json = response.choices[0].message['content'].strip()
            debug_print("Parsed response using newer format")
        except (AttributeError, KeyError, TypeError) as e:
            debug_print(f"Error parsing response with newer format: {str(e)}")
            try:
                # Try legacy format
                headers_json = response.choices[0].text.strip()
                debug_print("Parsed response using legacy format")
            except Exception as e2:
                debug_print(f"Error parsing response with legacy format: {str(e2)}")
                # Last resort - try direct access to dict
                headers_json = response['choices'][0]['message']['content'].strip()
        
        debug_print(f"Raw response: {headers_json}")
        
        # Handle cases where the response might include markdown or explanations
        if "```json" in headers_json:
            debug_print("Found JSON code block with json tag")
            headers_json = headers_json.split("```json")[1].split("```")[0].strip()
        elif "```" in headers_json:
            debug_print("Found generic code block")
            headers_json = headers_json.split("```")[1].split("```")[0].strip()
        
        # Try to extract just the JSON array if there's surrounding text
        headers_match = re.search(r'\[\s*"[^"]*"(?:\s*,\s*"[^"]*")*\s*\]', headers_json)
        if headers_match:
            debug_print("Extracted JSON array from response text")
            headers_json = headers_match.group(0)
        
        debug_print(f"Processed JSON: {headers_json}")
        
        try:
            headers = json.loads(headers_json)
            debug_print(f"Successfully parsed JSON: {headers}")
            return headers
        except json.JSONDecodeError as e:
            debug_print(f"JSON parse error: {str(e)}")
            # Try one more cleanup attempt - sometimes quotes are wrong
            clean_json = headers_json.replace("'", '"')
            debug_print(f"Attempting with cleaned JSON: {clean_json}")
            headers = json.loads(clean_json)
            return headers
            
    except Exception as e:
        st.error(f"Error identifying headers with OpenAI: {str(e)}")
        debug_print(f"Full exception: {traceback.format_exc()}")
        st.error(f"Response was: {headers_json if 'headers_json' in locals() else 'No response'}")
        return []

def process_document(docx_file, headers):
    """Process the document to extract text and bold words using the given headers"""
    # This function implements the core functionality from working2.py
    try:
        debug_print(f"Processing document with {len(headers)} headers: {headers}")
        doc = Document(docx_file)
        
        formatted_text = []
        extracted_words = []
        
        current_section = None
        current_subsection = None
        
        section_counter = 0
        
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
                
            text = para.text.strip()
            
            # Check if paragraph contains any headers
            header_match = False
            matching_header = None
            
            for keyword in headers:
                if keyword in text:
                    header_match = True
                    matching_header = keyword
                    break
                    
            if header_match:
                debug_print(f"Found header match: {matching_header} in paragraph {i}")
                section_counter += 1
                current_section = str(section_counter)
                
                processed_text = ""
                for run in para.runs:
                    if run.bold:
                        processed_text += f"*{run.text}*"
                    else:
                        processed_text += run.text
                
                formatted_text.append(f"{current_section}. {processed_text}")
                continue
            
            subsection_match = re.match(r'^([a-z])\.\s+(.*)', text)
            
            if subsection_match:
                subsection_letter = subsection_match.group(1)
                subsection_content = subsection_match.group(2)
                
                processed_text = ""
                for run in para.runs:
                    if run.bold:
                        processed_text += f"*{run.text}*"
                    else:
                        processed_text += run.text
                
                letter_part = processed_text[:2]
                content_part = processed_text[2:].lstrip()
                
                formatted_text.append(f"   {subsection_letter}. {content_part}")
                continue
            
            if current_section and len(formatted_text) > 0 and formatted_text[-1].startswith(f"{current_section}."):
                subsection_letter = "a"
                debug_print(f"Adding first subsection to section {current_section}")
                
                processed_text = ""
                for run in para.runs:
                    if run.bold:
                        processed_text += f"*{run.text}*"
                    else:
                        processed_text += run.text
                
                formatted_text.append(f"   {subsection_letter}. {processed_text}")
                current_subsection = subsection_letter
                continue
                
            if current_section and len(formatted_text) > 0:
                last_line = formatted_text[-1]
                if last_line.startswith("   "):
                    last_letter_match = re.match(r'\s+([a-z])\.\s+', last_line)
                    if last_letter_match:
                        last_letter = last_letter_match.group(1)
                        next_letter = chr(ord(last_letter) + 1)
                        debug_print(f"Adding sequential subsection {next_letter} to section {current_section}")
                        
                        processed_text = ""
                        for run in para.runs:
                            if run.bold:
                                processed_text += f"*{run.text}*"
                            else:
                                processed_text += run.text
                        
                        formatted_text.append(f"   {next_letter}. {processed_text}")
                        current_subsection = next_letter
                        continue
            
            processed_text = ""
            for run in para.runs:
                if run.bold:
                    processed_text += f"*{run.text}*"
                else:
                    processed_text += run.text
            
            formatted_text.append(f"      {processed_text}")
        
        debug_print(f"Initial document processing complete. Formatted {len(formatted_text)} paragraphs.")
        
        # Extract bold words
        bold_count = 0
        for line in formatted_text:
            if not line.strip():
                continue
            
            section_match = re.match(r'^(\d+)\.\s+', line)
            if section_match:
                current_section = section_match.group(1)
                current_subsection = None
                continue
            
            subsection_match = re.match(r'^\s*([a-z])\.\s+', line)
            if subsection_match:
                current_subsection = subsection_match.group(1)
            
            if current_section and current_subsection:
                bold_matches = re.findall(r'\*(.*?)\*', line)
                
                for bold_word in bold_matches:
                    if bold_word.strip():
                        section_id = f"{current_section}{current_subsection}"
                        entry = {
                            "section": current_section,
                            "subsection": current_subsection,
                            "section_id": section_id,
                            "text": bold_word.strip()
                        }
                        extracted_words.append(entry)
                        bold_count += 1
        
        debug_print(f"Bold word extraction complete. Found {bold_count} bold terms.")
        return formatted_text, extracted_words
        
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        debug_print(f"Full exception: {traceback.format_exc()}")
        return [], []

def main():
    st.title("Document Bold Text Extractor")
    
    # Add refresh button at the top
    if st.button("Start Over"):
        st.experimental_rerun()
    
    st.markdown("Upload a Word document to extract bold text organized by sections")
    
    # Debug controls
    with st.expander("Debug Settings"):
        global debug_mode
        debug_mode = st.checkbox("Enable Debug Mode", value=False)
        if st.button("Test OpenAI Connection"):
            try:
                if not openai.api_key:
                    st.error("API key not set")
                else:
                    debug_print("Testing OpenAI connection...")
                    try:
                        # Try completion API
                        response = openai.Completion.create(
                            engine="gpt-3.5-turbo-instruct",
                            prompt="Say hello",
                            max_tokens=5
                        )
                        st.success(f"OpenAI API Completion test successful: {response.choices[0].text}")
                    except Exception as e1:
                        debug_print(f"Completion API test failed: {str(e1)}")
                        try:
                            # Try chat completion API
                            response = openai.ChatCompletion.create(
                                model="gpt-3.5-turbo",
                                messages=[{"role": "user", "content": "Say hello"}]
                            )
                            st.success(f"OpenAI ChatCompletion API test successful")
                        except Exception as e2:
                            st.error(f"ChatCompletion API test failed: {str(e2)}")
            except Exception as e:
                st.error(f"Error testing API: {str(e)}")
    
    uploaded_file = st.file_uploader("Choose a Word document", type="docx")
    
    if uploaded_file is not None:
        debug_print(f"File uploaded: {uploaded_file.name}")
        
        # Save the uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(uploaded_file.getvalue())
            temp_file_path = tmp.name
            debug_print(f"Saved upload to temporary file: {temp_file_path}")
        
        try:
            with st.spinner("Processing document..."):
                # Convert document to text
                raw_text = convert_docx_to_raw_text(temp_file_path)
                debug_print(f"Converted document to {len(raw_text)} characters of text")
                
                # Identify headers with OpenAI
                headers = identify_headers_with_openai(raw_text)
                
                if not headers:
                    st.warning("No headers were identified. Please try a different document.")
                    return
                    
                debug_print(f"Identified {len(headers)} headers: {headers}")
                
                # Process document and extract bold words
                formatted_text, extracted_words = process_document(temp_file_path, headers)
                
                if not extracted_words:
                    st.warning("No bold words were found in the document.")
                    return
                
                # Format the bold words as a simple list grouped by section
                output_text = []
                sections = {}
                
                # Group words by section for display
                for word in extracted_words:
                    section_id = word["section"]
                    if section_id not in sections:
                        sections[section_id] = {}
                        
                    subsection_id = word["subsection"]
                    if subsection_id not in sections[section_id]:
                        sections[section_id][subsection_id] = []
                        
                    sections[section_id][subsection_id].append(word["text"])
                
                # Create a simple text representation
                for section_id in sorted(sections.keys()):
                    output_text.append(f"SECTION {section_id}")
                    for subsection_id in sorted(sections[section_id].keys()):
                        output_text.append(f"  {section_id}{subsection_id}:")
                        words = sections[section_id][subsection_id]
                        for word in words:
                            output_text.append(f"    • {word}")
                        output_text.append("")
                
                # Display the result in a text box
                st.success(f"Successfully extracted {len(extracted_words)} bold words")
                st.text_area("Extracted Bold Words", "\n".join(output_text), height=400)
                
                # Create simplified format for download (like zbold.txt)
                simplified_output = []
                for section_id in sorted(sections.keys()):
                    for subsection_id in sorted(sections[section_id].keys()):
                        section_subsection = f"{section_id}{subsection_id}"
                        words = sections[section_id][subsection_id]
                        for word in words:
                            simplified_output.append(f"{section_subsection}: {word}")
                
                # Add a download button for the text in simplified format
                st.download_button(
                    label="Download as Text",
                    data="\n".join(simplified_output),
                    file_name="extracted_bold_words.txt",
                    mime="text/plain"
                )
                
                # Add refresh button at the bottom
                st.write("")
                if st.button("Process Another Document"):
                    st.experimental_rerun()
                
        except Exception as e:
            st.error(f"Error processing document: {str(e)}")
            debug_print(f"Full exception: {traceback.format_exc()}")
        finally:
            # Clean up the temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                debug_print(f"Cleaned up temporary file: {temp_file_path}")

if __name__ == "__main__":
    main() 